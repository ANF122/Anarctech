import tkinter as tk
from tkinter import filedialog, messagebox, colorchooser, simpledialog
from tkinter import ttk
from threading import Thread
import time
import os
import shutil
from docx import Document
from spellchecker import SpellChecker
from PIL import Image, ImageTk
class FuturistAnarchistApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Anartech")
        self.root.geometry("1200x800")
        self.root.configure(bg="#000000")  # Noir pour le thème anarchiste
        self.root.iconbitmap("C:/Users/pirat/Anarchist.ico")
        # Create the installation menu
        self.show_installation_menu()

    def show_installation_menu(self):
        """Show the installation menu."""
        self.installation_window = tk.Toplevel(self.root)
        self.installation_window.title("Installation")
        self.installation_window.geometry("400x300")
        self.installation_window.configure(bg="#000000")  # Noir pour le thème anarchiste
        self.installation_window.iconbitmap("C:/Users/pirat/Anarchist.ico")
        tk.Label(self.installation_window, text="Bienvenue dans le programme d'installation", 
                 bg="#000000", fg="#FFFFFF", font=("Arial", 14)).pack(pady=10)

        tk.Button(self.installation_window, text="Commencer l'Installation", command=self.install_app,
                  bg="#333333", fg="#FFFFFF", relief=tk.FLAT).pack(pady=10)

        tk.Button(self.installation_window, text="Quitter", command=self.root.quit,
                  bg="#333333", fg="#FFFFFF", relief=tk.FLAT).pack(pady=10)

    def install_app(self):
        """Perform the installation process."""
        self.installation_window.withdraw()  # Hide the installation window
        self.setup_application()

    def setup_application(self):
        """Set up the application."""
        # Simulate setup process
        self.setup_window = tk.Toplevel(self.root)
        self.setup_window.title("Configuration")
        self.setup_window.geometry("400x300")
        self.setup_window.configure(bg="#000000")  # Noir pour le thème anarchiste
        self.setup_window.iconbitmap("C:/Users/pirat/Anarchist.ico")
        tk.Label(self.setup_window, text="Configuration de l'application en cours...",
                 bg="#000000", fg="#FFFFFF", font=("Arial", 14)).pack(pady=20)

        self.progress = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(self.setup_window, variable=self.progress, maximum=100)
        self.progress_bar.pack(pady=20, fill=tk.X, padx=20)

        self.status_label = tk.Label(self.setup_window, text="Préparation...", bg="#000000", fg="#FFFFFF")
        self.status_label.pack(pady=10)

        # Simulate setup process
        self.install_thread = Thread(target=self.run_setup_process)
        self.install_thread.start()

    def run_setup_process(self):
        """Simulate a setup process with progress updates."""
        for i in range(101):
            time.sleep(0.05)
            self.root.after(0, self.update_progress, i)
        self.root.after(0, self.setup_complete)

    def update_progress(self, value):
        """Update progress bar."""
        self.progress.set(value)
        self.status_label.config(text=f"Progression: {value}%")

    def setup_complete(self):
        """Finish the setup process."""
        self.setup_window.destroy()
        self.show_main_app()

    def show_main_app(self):
        """Show the main application interface."""
        # Create the menu
        self.menu = tk.Menu(self.root, bg="#111111", fg="#FFFFFF")
        self.root.config(menu=self.menu)
        self.root.iconbitmap("C:/Users/pirat/Anarchist.ico")  
        self.file_menu = tk.Menu(self.menu, tearoff=0, bg="#222222", fg="#FFFFFF")
        self.menu.add_cascade(label="Fichier", menu=self.file_menu)
        self.file_menu.add_command(label="Ouvrir", command=self.open_file)
        self.file_menu.add_command(label="Sauvegarder", command=self.save_file)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Quitter", command=self.root.quit)

        self.edit_menu = tk.Menu(self.menu, tearoff=0, bg="#222222", fg="#FFFFFF")
        self.menu.add_cascade(label="Édition", menu=self.edit_menu)
        self.edit_menu.add_command(label="Couper", command=lambda: self.text_area.event_generate("<<Cut>>"))
        self.edit_menu.add_command(label="Copier", command=lambda: self.text_area.event_generate("<<Copy>>"))
        self.edit_menu.add_command(label="Coller", command=lambda: self.text_area.event_generate("<<Paste>>"))
        self.edit_menu.add_command(label="Rechercher", command=self.find_text)

        self.format_menu = tk.Menu(self.menu, tearoff=0, bg="#222222", fg="#FFFFFF")
        self.menu.add_cascade(label="Format", menu=self.format_menu)
        self.format_menu.add_command(label="Police", command=self.choose_font)
        self.format_menu.add_command(label="Couleur du Texte", command=self.choose_color)
        self.format_menu.add_command(label="Changer Couleur de Fond", command=self.change_background_color)
        self.format_menu.add_command(label="Alignement Gauche", command=lambda: self.set_text_alignment('left'))
        self.format_menu.add_command(label="Alignement Centre", command=lambda: self.set_text_alignment('center'))
        self.format_menu.add_command(label="Alignement Droit", command=lambda: self.set_text_alignment('right'))

        # Create the toolbar
        self.toolbar = tk.Frame(self.root, bg="#111111")
        self.toolbar.pack(side=tk.TOP, fill=tk.X)

        self.bold_button = tk.Button(self.toolbar, text="Gras", command=self.toggle_bold, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.bold_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.italic_button = tk.Button(self.toolbar, text="Italique", command=self.toggle_italic, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.italic_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.underline_button = tk.Button(self.toolbar, text="Souligné", command=self.toggle_underline, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.underline_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.font_size = tk.StringVar(value="12")
        self.font_size_menu = tk.OptionMenu(self.toolbar, self.font_size, *[str(i) for i in range(8, 73, 2)], command=self.change_font_size)
        self.font_size_menu.pack(side=tk.LEFT, padx=2, pady=2)

        self.font_family = tk.StringVar(value="Arial")
        self.font_family_menu = tk.OptionMenu(self.toolbar, self.font_family, "Arial", "Courier", "Helvetica", "Times New Roman", command=self.change_font_family)
        self.font_family_menu.pack(side=tk.LEFT, padx=2, pady=2)

        # Create the text area
        self.text_area = tk.Text(self.root, wrap='word', bg="#1e1e1e", fg="#FFFFFF", font=("Arial", 12))
        self.text_area.pack(expand=1, fill='both')

        # Create buttons for additional features
        self.button_frame = tk.Frame(self.root, bg="#000000")
        self.button_frame.pack(side=tk.BOTTOM, fill=tk.X)

        self.add_button = tk.Button(self.button_frame, text="Ajouter Utilisateur", command=self.add_user, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.add_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.remove_button = tk.Button(self.button_frame, text="Supprimer Utilisateur", command=self.remove_user, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.remove_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.show_citations_button = tk.Button(self.button_frame, text="Afficher Citation", command=self.show_citations, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.show_citations_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.start_notifications_button = tk.Button(self.button_frame, text="Démarrer Notifications", command=self.start_notifications, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.start_notifications_button.pack(side=tk.LEFT, padx=2, pady=2)

        self.stop_notifications_button = tk.Button(self.button_frame, text="Arrêter Notifications", command=self.stop_notifications, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.stop_notifications_button.pack(side=tk.LEFT, padx=2, pady=2)
       
       # Bouton de vérification de l'orthographe
        self.check_spelling_button = tk.Button(self.button_frame, text="Vérifier Orthographe", command=self.check_spelling, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.check_spelling_button.pack(side=tk.LEFT, padx=2, pady=2)
        
        self.insert_image_button = tk.Button(self.button_frame, text="Insérer Image", command=self.insert_image, bg="#333333", fg="#FFFFFF", relief=tk.FLAT)
        self.insert_image_button.pack(side=tk.LEFT, padx=2, pady=2)  

        # Create a listbox to manage users
        self.users = []
        self.users_listbox = tk.Listbox(self.root, bg="#333333", fg="#FFFFFF", font=("Arial", 12))
        self.users_listbox.pack(side=tk.LEFT, fill=tk.Y)
        self.text_editor_window = None
        self.notification_thread = None
        self.notification_thread_running = False



    
    def set_text_alignment(self, alignment):
        """Set the alignment of the entire text."""
        # Configure the tag for alignment
        self.text_area.tag_configure("align", justify=alignment)

        # Apply the alignment tag to the entire text
        self.text_area.tag_remove("align", "1.0", "end")  # Remove previous alignment tags
        self.text_area.tag_add("align", "1.0", "end")  # Apply new alignment to all text
    def check_spelling(self):
        """Vérifie l'orthographe du texte dans la zone de texte et demande confirmation pour chaque correction."""
        spell = SpellChecker()
        content = self.text_area.get(1.0, tk.END).strip()
        words = content.split()

        misspelled = spell.unknown(words)

        if misspelled:
            for word in misspelled:
                # Obtenir les suggestions de correction
                suggestions = list(spell.candidates(word))
                if suggestions:
                    # Convertir les suggestions en chaîne de caractères
                    suggestion_text = ', '.join(suggestions)
                    # Demander confirmation à l'utilisateur
                    response = messagebox.askyesno(
                        "Correction orthographique",
                        f"Le mot '{word}' est mal orthographié. Souhaitez-vous le remplacer par l'une des suggestions suivantes : {suggestion_text} ?"
                    )
                    if response:
                        # Remplacer le mot par la première suggestion
                        content = content.replace(word, suggestions[0])
        
            # Met à jour le texte dans la zone de texte
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(tk.END, content)
        else:
            messagebox.showinfo("Vérification d'orthographe", "Aucun mot mal orthographié trouvé.")


 
    def change_background_color(self):
        """Allow the user to choose a background color for the text area."""
        color_code = colorchooser.askcolor(title="Choisir Couleur de Fond")[1]
        if color_code:
            self.text_area.configure(bg=color_code)
    def add_user(self):
        """Add a new user to the list."""
        user_name = simpledialog.askstring("Ajouter Utilisateur", "Entrez le nom de l'utilisateur:")
        if user_name:
            self.users.append(user_name)
            self.users_listbox.insert(tk.END, user_name)

    def remove_user(self):
        """Remove a selected user from the list."""
        selected_index = self.users_listbox.curselection()
        if selected_index:
            self.users_listbox.delete(selected_index)
            del self.users[selected_index[0]]

    def show_citations(self):
        """Show a list of citations."""
        citations = [
            "Mikhail Bakounine : 'La liberté sans socialisme est un privilège, l'inégalité; le socialisme sans liberté est une servitude.'",
        "Emma Goldman : 'Si je ne peux danser, ce n'est pas ma révolution.'",
        "Pierre Kropotkine : 'L'anarchisme est un principe de vie, une philosophie qui est pratiquée tous les jours, par tous ceux qui cherchent à vivre librement et en harmonie avec les autres.'",
        "Noam Chomsky : 'L'anarchisme est un système de pensée qui ne repose pas sur la force, mais sur le consentement.'",
        "Lucy Parsons : 'Il n'y a pas de paix dans l'esclavage.'",
        "Alexander Berkman : 'La lutte pour la liberté est la lutte pour la vie.'",
        "David Graeber : 'Le véritable pouvoir est celui qui ne se voit pas.'",
        "Pierre-Joseph Proudhon : 'La propriété, c'est le vol.'"
        ]
        citations_text = "\n".join(citations)
        messagebox.showinfo("Citations", citations_text)

    def start_notifications(self):
        """Start a thread to handle notifications."""
        if not self.notification_thread_running:
            self.notification_thread_running = True
            self.notification_thread = Thread(target=self.notification_loop)
            self.notification_thread.start()

    def stop_notifications(self):
        """Stop the notification thread."""
        self.notification_thread_running = False
        if self.notification_thread:
            self.notification_thread.join()

    def notification_loop(self):
        """Loop to generate notifications."""
        while self.notification_thread_running:
            self.root.after(0, lambda: messagebox.showinfo("Notification", "Pierre-Joseph Proudhon : 'La propriété, c'est le vol"))
            time.sleep(10)  # Wait for 10 seconds before showing another notification

    def open_file(self):
        """Open a file dialog to select and open a file."""
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            if file_path.endswith(".docx"):
                self.open_docx_file(file_path)
            else:
                self.open_txt_file(file_path)

    def open_txt_file(self, file_path):
        """Open a TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(tk.END, content)

    def open_docx_file(self, file_path):
        """Open a DOCX file."""
        doc = Document(file_path)
        content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, content)

    def save_file(self):
        """Save the content of the text area to a file."""
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            if file_path.endswith(".docx"):
                self.save_docx_file(file_path)
            else:
                self.save_txt_file(file_path)

    def save_txt_file(self, file_path):
        """Save the content as a TXT file."""
        with open(file_path, 'w', encoding='utf-8') as file:
            content = self.text_area.get(1.0, tk.END)
            file.write(content)

    def save_docx_file(self, file_path):
        """Save the content as a DOCX file."""
        doc = Document()
        content = self.text_area.get(1.0, tk.END).strip()
        doc.add_paragraph(content)
        doc.save(file_path)

    def choose_font(self):
        """Allow the user to choose a font for the text."""
        font_name = simpledialog.askstring("Choisir Police", "Entrez le nom de la police:")
        if font_name:
            self.text_area.configure(font=(font_name, self.font_size.get()))

    def choose_color(self):
        """Allow the user to choose a color for the text."""
        color_code = colorchooser.askcolor(title="Choisir Couleur")[1]
        if color_code:
            self.text_area.configure(fg=color_code)

    def toggle_bold(self):
        """Toggle bold formatting."""
        current_tags = self.text_area.tag_names("sel.first")
        if "bold" in current_tags:
            self.text_area.tag_remove("bold", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("bold", "sel.first", "sel.last")
            self.text_area.tag_configure("bold", font=("Arial", 12, "bold"))

    def toggle_italic(self):
        """Toggle italic formatting."""
        current_tags = self.text_area.tag_names("sel.first")
        if "italic" in current_tags:
            self.text_area.tag_remove("italic", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("italic", "sel.first", "sel.last")
            self.text_area.tag_configure("italic", font=("Arial", 12, "italic"))

    def toggle_underline(self):
        """Toggle underline formatting."""
        current_tags = self.text_area.tag_names("sel.first")
        if "underline" in current_tags:
            self.text_area.tag_remove("underline", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("underline", "sel.first", "sel.last")
            self.text_area.tag_configure("underline", font=("Arial", 12, "underline"))

    def change_font_size(self, size):
        """Change the font size."""
        self.text_area.configure(font=(self.font_family.get(), size))

    def change_font_family(self, family):
        """Change the font family."""
        self.text_area.configure(font=(family, self.font_size.get()))


    def find_text(self):
        """Find text in the text area."""
        search_term = simpledialog.askstring("Recherche", "Entrez le texte à rechercher:")
        if search_term:
            content = self.text_area.get(1.0, tk.END)
            start_index = content.find(search_term)
            if start_index != -1:
                start_index = self.text_area.index(f"1.0 + {start_index}c")
                end_index = self.text_area.index(f"{start_index} + {len(search_term)}c")
                self.text_area.tag_add("highlight", start_index, end_index)
                self.text_area.tag_configure("highlight", background="yellow")
            else:
                messagebox.showinfo("Recherche", "Texte non trouvé.")

if __name__ == "__main__":
    root = tk.Tk()
    app = FuturistAnarchistApp(root)
    root.mainloop()